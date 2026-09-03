import ast
import base64
import hashlib
import hmac
import json
import os
import re
import secrets
import time
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import httpx
from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel, Field
from sqlalchemy import DateTime, ForeignKey, Integer, String, Text, create_engine, inspect, text
from sqlalchemy.orm import DeclarativeBase, Mapped, mapped_column, relationship, sessionmaker

DATABASE_URL = os.getenv("DATABASE_URL", "mysql+pymysql://yours_ai:yours_ai@localhost:3306/yours_ai?charset=utf8")
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://127.0.0.1:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen3:4b")
OLLAMA_KEEP_ALIVE = os.getenv("OLLAMA_KEEP_ALIVE", "10m")
IMAGE_GENERATION_PROVIDER = os.getenv("IMAGE_GENERATION_PROVIDER", "")
IMAGE_GENERATION_URL = os.getenv("IMAGE_GENERATION_URL", "")
TOKEN_SECRET = os.getenv("TOKEN_SECRET", "change-this-local-secret")
CORS_ORIGINS = [x.strip() for x in os.getenv("CORS_ORIGINS", "http://127.0.0.1:8080,http://localhost:8080").split(",") if x.strip()]
try:
    MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_SIZE_MB", "10")) * 1024 * 1024
except ValueError:
    MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_CONTEXT_MESSAGES = 20
MODEL_CACHE_TTL_SECONDS = 30
_model_cache: tuple[float, list[str]] = (0, [])
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
IMAGE_DIR = UPLOAD_DIR / "images"
IMAGE_DIR.mkdir(exist_ok=True)
GENERATED_DIR = Path("generated")
GENERATED_DIR.mkdir(exist_ok=True)

def mysql_safe_text(value: str) -> str:
    """Keep MySQL's three-byte utf8 columns from receiving four-byte Unicode."""
    return "".join(character if ord(character) <= 0xFFFF else "?" for character in value)

engine = create_engine(DATABASE_URL, pool_pre_ping=True)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)

class Base(DeclarativeBase):
    pass

class User(Base):
    __tablename__ = "users"
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    name: Mapped[str] = mapped_column(String(120))
    email: Mapped[str] = mapped_column(String(255), unique=True, index=True)
    password_hash: Mapped[str] = mapped_column(String(255))
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    conversations = relationship("Conversation", back_populates="user", cascade="all, delete-orphan")

class Conversation(Base):
    __tablename__ = "conversations"
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), index=True, nullable=True)
    title: Mapped[str] = mapped_column(String(255), default="New chat")
    model: Mapped[str] = mapped_column(String(120), default=OLLAMA_MODEL)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    user = relationship("User", back_populates="conversations")
    messages = relationship("Message", back_populates="conversation", cascade="all, delete-orphan")

class Message(Base):
    __tablename__ = "messages"
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    conversation_id: Mapped[int] = mapped_column(ForeignKey("conversations.id"), index=True)
    role: Mapped[str] = mapped_column(String(30))
    content: Mapped[str] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    conversation = relationship("Conversation", back_populates="messages")

class Document(Base):
    __tablename__ = "documents"
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), index=True, nullable=True)
    conversation_id: Mapped[int | None] = mapped_column(ForeignKey("conversations.id"), nullable=True)
    filename: Mapped[str] = mapped_column(String(255))
    content: Mapped[str] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

class DocumentChunk(Base):
    __tablename__ = "document_chunks"
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    document_id: Mapped[int] = mapped_column(ForeignKey("documents.id"), index=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), index=True)
    content: Mapped[str] = mapped_column(Text)

class UserSetting(Base):
    __tablename__ = "user_settings"
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), primary_key=True)
    theme: Mapped[str] = mapped_column(String(20), default="dark")
    web_search: Mapped[bool] = mapped_column(default=False)
    coding_mode: Mapped[bool] = mapped_column(default=False)

class Attachment(Base):
    __tablename__ = "attachments"
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), index=True)
    conversation_id: Mapped[int | None] = mapped_column(ForeignKey("conversations.id"), nullable=True)
    filename: Mapped[str] = mapped_column(String(255))
    mime_type: Mapped[str] = mapped_column(String(120))
    size: Mapped[int] = mapped_column(Integer)
    storage_path: Mapped[str] = mapped_column(String(500))
    original_filename: Mapped[str | None] = mapped_column(String(255), nullable=True)
    stored_filename: Mapped[str | None] = mapped_column(String(255), nullable=True)
    image_format: Mapped[str | None] = mapped_column(String(30), nullable=True)
    analysis_path: Mapped[str | None] = mapped_column(String(500), nullable=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

class GeneratedFile(Base):
    __tablename__ = "generated_files"
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), index=True)
    filename: Mapped[str] = mapped_column(String(255))
    language: Mapped[str] = mapped_column(String(30))
    storage_path: Mapped[str] = mapped_column(String(500))
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

Base.metadata.create_all(engine)

def add_legacy_columns() -> None:
    table_columns = {
        "conversations": {"user_id": "INTEGER NULL", "model": "VARCHAR(120) NOT NULL DEFAULT 'qwen3:4b'"},
        "documents": {"user_id": "INTEGER NULL", "conversation_id": "INTEGER NULL"},
        "attachments": {
            "original_filename": "VARCHAR(255) NULL",
            "stored_filename": "VARCHAR(255) NULL",
            "image_format": "VARCHAR(30) NULL",
            "analysis_path": "VARCHAR(500) NULL",
        },
    }
    inspector = inspect(engine)
    with engine.begin() as connection:
        for table_name, columns in table_columns.items():
            existing = {column["name"] for column in inspector.get_columns(table_name)}
            for column_name, definition in columns.items():
                if column_name not in existing:
                    connection.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {definition}"))

add_legacy_columns()

def password_hash(password: str) -> str:
    salt = secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 310000)
    return f"pbkdf2_sha256$310000${salt.hex()}${digest.hex()}"

def password_valid(password: str, stored: str) -> bool:
    try:
        algorithm, rounds, salt, digest = stored.split("$")
        candidate = hashlib.pbkdf2_hmac("sha256", password.encode(), bytes.fromhex(salt), int(rounds)).hex()
        return algorithm == "pbkdf2_sha256" and hmac.compare_digest(candidate, digest)
    except (ValueError, TypeError):
        return False

def make_token(user_id: int) -> str:
    payload = f"{user_id}:{int(datetime.now(timezone.utc).timestamp())}"
    signature = hmac.new(TOKEN_SECRET.encode(), payload.encode(), hashlib.sha256).hexdigest()
    return f"{payload}:{signature}"

def current_user(credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer())) -> User:
    try:
        user_id, issued, signature = credentials.credentials.split(":")
        payload = f"{user_id}:{issued}"
        expected = hmac.new(TOKEN_SECRET.encode(), payload.encode(), hashlib.sha256).hexdigest()
        if not hmac.compare_digest(signature, expected) or datetime.now(timezone.utc).timestamp() - int(issued) > 86400 * 7:
            raise ValueError
    except (ValueError, TypeError):
        raise HTTPException(401, "Invalid or expired authentication token")
    with SessionLocal() as db:
        user = db.get(User, int(user_id))
        if not user:
            raise HTTPException(401, "User not found")
        db.expunge(user)
        return user

def owned_conversation(db, conversation_id: int, user_id: int) -> Conversation:
    row = db.query(Conversation).filter(Conversation.id == conversation_id, Conversation.user_id == user_id).first()
    if not row:
        raise HTTPException(404, "Conversation not found")
    return row

def split_chunks(content: str) -> list[str]:
    words = content.split()
    return [" ".join(words[i:i + 180]) for i in range(0, len(words), 150) if words[i:i + 180]]

def relevant_chunks(db, user_id: int, query: str, limit: int = 4) -> list[tuple[str, str]]:
    terms = set(re.findall(r"[a-z0-9]{3,}", query.lower()))
    scored = []
    for chunk, filename in db.query(DocumentChunk.content, Document.filename).join(Document, Document.id == DocumentChunk.document_id).filter(DocumentChunk.user_id == user_id):
        score = sum(term in chunk.lower() for term in terms)
        if score:
            scored.append((score, chunk, filename))
    return [(chunk, filename) for score, chunk, filename in sorted(scored, reverse=True)[:limit]]

def calculator(expression: str) -> str | None:
    candidate = expression.strip().replace("calculate", "", 1).strip()
    if not re.fullmatch(r"[0-9+*/().%\-\s]+", candidate):
        return None
    try:
        tree = ast.parse(candidate, mode="eval")
        if any(isinstance(node, (ast.Name, ast.Call, ast.Attribute, ast.Subscript)) for node in ast.walk(tree)):
            return None
        return str(eval(compile(tree, "<calculator>", "eval"), {"__builtins__": {}}, {}))
    except (SyntaxError, ZeroDivisionError, ValueError):
        return None

def safe_filename(filename: str, fallback: str) -> str:
    cleaned = Path(filename or fallback).name
    return cleaned if cleaned not in {"", ".", ".."} else fallback

def allowed_image(filename: str, content_type: str) -> bool:
    return Path(filename).suffix.lower() in {".jpg", ".jpeg", ".png", ".webp", ".gif"} and content_type in {"image/jpeg", "image/png", "image/webp", "image/gif"}

def image_signature_matches(data: bytes, content_type: str) -> bool:
    signatures = {
        "image/jpeg": data.startswith(b"\xff\xd8\xff"),
        "image/png": data.startswith(b"\x89PNG\r\n\x1a\n"),
        "image/gif": data.startswith((b"GIF87a", b"GIF89a")),
        "image/webp": data.startswith(b"RIFF") and data[8:12] == b"WEBP",
    }
    return signatures.get(content_type, False)

async def ollama_capabilities(model: str) -> list[str]:
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.post(f"{OLLAMA_URL}/api/show", json={"name": model})
            if response.status_code == 404:
                return []
            response.raise_for_status()
            return response.json().get("capabilities", [])
    except httpx.HTTPError as exc:
        raise HTTPException(503, f"Ollama unavailable: {exc}") from exc

async def search_results(query: str) -> list[dict[str, str]]:
    try:
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
            response = await client.post("https://html.duckduckgo.com/html/", data={"q": query[:300]}, headers={"User-Agent": "Yours AI/2.0"})
            response.raise_for_status()
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(response.text, "html.parser")
        results = []
        for item in soup.select(".result")[:5]:
            link = item.select_one(".result__a")
            snippet = item.select_one(".result__snippet")
            if link:
                results.append({"title": link.get_text(" ", strip=True), "url": link.get("href", ""), "snippet": snippet.get_text(" ", strip=True) if snippet else ""})
        return results
    except (httpx.HTTPError, ValueError):
        return []

app = FastAPI(title="Yours AI API", version="2.0.0")
app.add_middleware(CORSMiddleware, allow_origins=CORS_ORIGINS, allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

class RegisterRequest(BaseModel):
    name: str = Field(min_length=1, max_length=120)
    email: str = Field(min_length=5, max_length=255)
    password: str = Field(min_length=8, max_length=128)

class LoginRequest(BaseModel):
    email: str
    password: str

class ChatRequest(BaseModel):
    message: str = Field(min_length=1, max_length=20000)
    conversation_id: int | None = None
    model: str | None = None
    system_prompt: str | None = None
    search_enabled: bool = False
    coding_mode: bool = False
    document_ids: list[int] = []
    attachment_ids: list[int] = []

@app.get("/")
def root():
    return {"app": "Yours AI", "status": "running"}

@app.get("/health")
def health():
    return {"status": "ok", "ollama_url": OLLAMA_URL, "configured_model": OLLAMA_MODEL}

@app.post("/auth/register")
def register(req: RegisterRequest):
    email = req.email.strip().lower()
    with SessionLocal() as db:
        if db.query(User).filter(User.email == email).first():
            raise HTTPException(409, "An account with that email already exists")
        user = User(name=mysql_safe_text(req.name.strip()), email=email, password_hash=password_hash(req.password))
        db.add(user)
        db.commit()
        db.refresh(user)
        return {"token": make_token(user.id), "user": {"id": user.id, "name": user.name, "email": user.email}}

@app.post("/auth/login")
def login(req: LoginRequest):
    with SessionLocal() as db:
        user = db.query(User).filter(User.email == req.email.strip().lower()).first()
        if not user or not password_valid(req.password, user.password_hash):
            raise HTTPException(401, "Invalid email or password")
        return {"token": make_token(user.id), "user": {"id": user.id, "name": user.name, "email": user.email}}

@app.post("/auth/logout")
def logout(_: User = Depends(current_user)):
    return {"message": "logged out"}

@app.get("/auth/me")
def me(user: User = Depends(current_user)):
    return {"id": user.id, "name": user.name, "email": user.email}

async def installed_models() -> list[str]:
    global _model_cache
    cached_at, cached_models = _model_cache
    if time.monotonic() - cached_at < MODEL_CACHE_TTL_SECONDS:
        return cached_models
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(f"{OLLAMA_URL}/api/tags")
            response.raise_for_status()
            models = [item["name"] for item in response.json().get("models", [])]
            _model_cache = (time.monotonic(), models)
            return models
    except httpx.HTTPError as exc:
        raise HTTPException(503, f"Ollama unavailable: {exc}")

@app.get("/models")
async def models(_: User = Depends(current_user)):
    return {"models": await installed_models(), "default": OLLAMA_MODEL}

@app.get("/models/capabilities")
async def model_capabilities(user: User = Depends(current_user)):
    del user
    names = await installed_models()
    return {"models": [{"name": name, "capabilities": await ollama_capabilities(name)} for name in names]}

@app.post("/upload/image")
async def upload_image(file: UploadFile = File(...), conversation_id: int | None = Form(None), user: User = Depends(current_user)):
    from services.image_service import inspect_image, supported_mime
    if not supported_mime(file.filename or "", file.content_type or ""):
        raise HTTPException(415, "Supported images are JPG, JPEG, PNG, WEBP, GIF, BMP, TIFF, SVG, HEIC, HEIF, ICO, and AVIF with matching MIME types")
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "Image exceeds the 10 MB limit")
    try:
        image_format, width, height = inspect_image(data, file.filename or "", file.content_type or "")
    except ValueError as exc:
        raise HTTPException(415, str(exc)) from exc
    with SessionLocal() as db:
        if conversation_id is not None:
            owned_conversation(db, conversation_id, user.id)
        stored_name = f"{uuid.uuid4().hex}{Path(file.filename or '.img').suffix.lower()}"
        path = IMAGE_DIR / stored_name
        path.write_bytes(data)
        attachment = Attachment(user_id=user.id, conversation_id=conversation_id, filename=safe_filename(file.filename or stored_name, stored_name), original_filename=safe_filename(file.filename or stored_name, stored_name), stored_filename=stored_name, image_format=image_format, mime_type=file.content_type, size=len(data), storage_path=str(path))
        db.add(attachment)
        db.commit()
        db.refresh(attachment)
        return {"id": attachment.id, "filename": attachment.filename, "mime_type": attachment.mime_type, "size": attachment.size, "format": attachment.image_format, "width": width, "height": height, "download_url": f"/files/attachments/{attachment.id}/download"}

@app.get("/files/attachments/{attachment_id}/download")
def download_attachment(attachment_id: int, user: User = Depends(current_user)):
    with SessionLocal() as db:
        attachment = db.query(Attachment).filter(Attachment.id == attachment_id, Attachment.user_id == user.id).first()
        if not attachment:
            raise HTTPException(404, "Attachment not found")
        path = Path(attachment.storage_path).resolve()
        if path.parent != IMAGE_DIR.resolve() or not path.is_file():
            raise HTTPException(404, "Attachment not found")
        return FileResponse(path, media_type=attachment.mime_type, filename=attachment.original_filename or attachment.filename)

@app.post("/upload")
async def upload_attachment(file: UploadFile = File(...), conversation_id: int | None = Form(None), user: User = Depends(current_user)):
    allowed = {".txt", ".md", ".pdf", ".docx", ".csv", ".json", ".py", ".js", ".html", ".css", ".sql", ".zip", ".exe"}
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in allowed:
        raise HTTPException(415, "This file type is not supported")
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File exceeds the 10 MB limit")
    with SessionLocal() as db:
        if conversation_id is not None:
            owned_conversation(db, conversation_id, user.id)
        folder = UPLOAD_DIR / ("executables" if suffix == ".exe" else "attachments")
        folder.mkdir(exist_ok=True)
        stored_name = f"{uuid.uuid4().hex}{suffix}"
        path = folder / stored_name
        path.write_bytes(data)
        attachment = Attachment(user_id=user.id, conversation_id=conversation_id, filename=safe_filename(file.filename or stored_name, stored_name), mime_type=file.content_type or "application/octet-stream", size=len(data), storage_path=str(path))
        db.add(attachment)
        db.commit()
        db.refresh(attachment)
        return {"id": attachment.id, "filename": attachment.filename, "mime_type": attachment.mime_type, "size": attachment.size, "executable": suffix == ".exe"}

@app.post("/vision/analyze")
async def vision_analyze(attachment_id: int | None = Form(None), attachment_ids: list[int] = Form([]), prompt: str = Form("Describe this image."), model: str | None = Form(None), conversation_id: int | None = Form(None), user: User = Depends(current_user)):
    from services.image_service import create_analysis_copy
    selected_model = model or OLLAMA_MODEL
    capabilities = await ollama_capabilities(selected_model)
    if "vision" not in capabilities:
        raise HTTPException(400, f"The selected model '{selected_model}' does not support image analysis. Select or install a vision-capable Ollama model.")
    ids = list(dict.fromkeys(attachment_ids or ([attachment_id] if attachment_id is not None else [])))
    if not ids or len(ids) > 4:
        raise HTTPException(400, "Provide between 1 and 4 image attachments")
    image_data_list = []
    converted_mime = "image/png"
    with SessionLocal() as db:
        attachments = db.query(Attachment).filter(Attachment.id.in_(ids), Attachment.user_id == user.id).all()
        if len(attachments) != len(ids) or any(not item.mime_type.startswith("image/") for item in attachments):
            raise HTTPException(404, "One or more image attachments were not found")
        for attachment in attachments:
            original_data = Path(attachment.storage_path).read_bytes()
            analysis_path = IMAGE_DIR / f"analysis_{attachment.id}.png"
            try:
                converted_path, converted_mime = create_analysis_copy(original_data, attachment.original_filename or attachment.filename, attachment.mime_type, analysis_path)
            except ValueError as exc:
                raise HTTPException(415, str(exc)) from exc
            attachment.analysis_path = str(converted_path)
            image_data_list.append(base64.b64encode(converted_path.read_bytes()).decode("ascii"))
        db.commit()
    payload = {"model": selected_model, "stream": False, "keep_alive": OLLAMA_KEEP_ALIVE, "messages": [{"role": "user", "content": prompt[:20000], "images": image_data_list}]}
    try:
        async with httpx.AsyncClient(timeout=300) as client:
            response = await client.post(f"{OLLAMA_URL}/api/chat", json=payload)
            response.raise_for_status()
            answer = response.json().get("message", {}).get("content", "").strip()
    except httpx.HTTPError as exc:
        raise HTTPException(502, f"Vision model unavailable: {exc}") from exc
    if not answer:
        raise HTTPException(502, "Vision model returned an empty response")
    with SessionLocal() as db:
        conversation = owned_conversation(db, conversation_id, user.id) if conversation_id else Conversation(user_id=user.id, title=mysql_safe_text(prompt[:60]), model=selected_model)
        if not conversation_id:
            db.add(conversation)
            db.flush()
        db.add_all([
            Message(conversation_id=conversation.id, role="user", content=mysql_safe_text(prompt)),
            Message(conversation_id=conversation.id, role="assistant", content=mysql_safe_text(answer)),
        ])
        db.commit()
        return {"response": answer, "model": selected_model, "attachment_id": ids[0], "attachment_ids": ids, "conversation_id": conversation.id, "processed_download_url": f"/files/attachments/{ids[0]}/analysis/download", "processed_mime_type": converted_mime}

@app.get("/files/attachments/{attachment_id}/analysis/download")
def download_analysis_copy(attachment_id: int, user: User = Depends(current_user)):
    with SessionLocal() as db:
        attachment = db.query(Attachment).filter(Attachment.id == attachment_id, Attachment.user_id == user.id).first()
        if not attachment or not attachment.analysis_path:
            raise HTTPException(404, "Processed image not found")
        path = Path(attachment.analysis_path).resolve()
        if path.parent != IMAGE_DIR.resolve() or not path.is_file():
            raise HTTPException(404, "Processed image not found")
        return FileResponse(path, media_type="image/png", filename=f"{Path(attachment.original_filename or attachment.filename).stem}_analysis.png")

@app.post("/image/generate")
async def image_generate(prompt: str = Form(...), user: User = Depends(current_user)):
    del user
    from services.image_generation import ImageGenerationError, generate_image
    try:
        data, media_type, extension = await generate_image(prompt[:20000])
    except ImageGenerationError as exc:
        raise HTTPException(503, str(exc)) from exc
    filename = f"{uuid.uuid4().hex}{extension}"
    path = GENERATED_DIR / filename
    path.write_bytes(data)
    return {"filename": filename, "mime_type": media_type, "download_url": f"/files/generated/{filename}"}

@app.get("/files/generated/{filename}")
def generated_file(filename: str, _: User = Depends(current_user)):
    path = GENERATED_DIR / Path(filename).name
    if not path.exists() or path.parent != GENERATED_DIR:
        raise HTTPException(404, "Generated file not found")
    return FileResponse(path)

class GeneratedCodeRequest(BaseModel):
    filename: str = Field(min_length=1, max_length=120)
    content: str = Field(max_length=500000)

@app.post("/code/files")
def create_code_file(req: GeneratedCodeRequest, user: User = Depends(current_user)):
    suffix = Path(req.filename).suffix.lower()
    allowed = {".py", ".js", ".html", ".css", ".sql", ".json", ".md", ".txt"}
    if suffix not in allowed:
        raise HTTPException(415, "Generated source file type is not allowed")
    safe_name = safe_filename(req.filename, f"generated{suffix}")
    path = GENERATED_DIR / f"{uuid.uuid4().hex}_{safe_name}"
    path.write_text(req.content, encoding="utf-8")
    with SessionLocal() as db:
        generated = GeneratedFile(user_id=user.id, filename=safe_name, language=suffix[1:], storage_path=str(path))
        db.add(generated)
        db.commit()
        db.refresh(generated)
        return {"id": generated.id, "filename": generated.filename, "language": generated.language, "download_url": f"/files/{generated.id}/download"}

@app.get("/files/{file_id}/download")
def download_generated_file(file_id: int, user: User = Depends(current_user)):
    with SessionLocal() as db:
        generated = db.query(GeneratedFile).filter(GeneratedFile.id == file_id, GeneratedFile.user_id == user.id).first()
        if not generated:
            raise HTTPException(404, "Generated file not found")
        return FileResponse(generated.storage_path, filename=generated.filename)

class ProjectRequest(BaseModel):
    files: dict[str, str] = Field(min_length=1, max_length=50)

@app.post("/code/project")
def create_project(req: ProjectRequest, user: User = Depends(current_user)):
    project_path = GENERATED_DIR / uuid.uuid4().hex
    project_path.mkdir()
    try:
        with zipfile.ZipFile(project_path / "project.zip", "w", zipfile.ZIP_DEFLATED) as archive:
            for filename, content in req.files.items():
                safe_path = Path(filename)
                if safe_path.is_absolute() or ".." in safe_path.parts or safe_path.suffix.lower() == ".exe":
                    raise HTTPException(400, "Unsafe project file path")
                if len(content) > 500000:
                    raise HTTPException(413, "Project file exceeds the 500 KB limit")
                archive.writestr(str(safe_path), content)
    except HTTPException:
        raise
    with SessionLocal() as db:
        generated = GeneratedFile(user_id=user.id, filename="project.zip", language="zip", storage_path=str(project_path / "project.zip"))
        db.add(generated)
        db.commit()
        db.refresh(generated)
        return {"id": generated.id, "filename": generated.filename, "files": list(req.files), "download_url": f"/files/{generated.id}/download"}

def system_prompt(req: ChatRequest, sources: list[tuple[str, str]]) -> str:
    parts = ["You are Yours AI, a helpful local assistant. Never claim to use a tool you did not use."]
    if req.coding_mode:
        parts.append("You are in coding assistant mode. Explain code safely and do not execute generated code.")
    if sources:
        parts.append("Retrieved document context (use only when relevant):\n" + "\n\n".join(f"[{name}] {chunk}" for chunk, name in sources))
    return "\n\n".join(parts)

async def ollama_stream(payload: dict):
    async with httpx.AsyncClient(timeout=300) as client:
        async with client.stream("POST", f"{OLLAMA_URL}/api/chat", json=payload) as response:
            if response.status_code != 200:
                detail = (await response.aread()).decode(errors="replace")[:1000]
                raise HTTPException(502, f"Ollama error: {detail}")
            async for line in response.aiter_lines():
                if line:
                    yield json.loads(line)

@app.post("/chat")
async def chat(req: ChatRequest, user: User = Depends(current_user)):
    request_started = time.perf_counter()
    print("[PERF] request received", flush=True)
    message = req.message.strip()
    model_started = time.perf_counter()
    available = await installed_models()
    print(f"[PERF] model validation: {(time.perf_counter() - model_started) * 1000:.1f} ms", flush=True)
    selected_model = req.model or OLLAMA_MODEL
    if selected_model not in available:
        raise HTTPException(400, f"Model '{selected_model}' is not installed. Available models: {', '.join(available) or 'none'}")
    database_started = time.perf_counter()
    with SessionLocal() as db:
        conversation = owned_conversation(db, req.conversation_id, user.id) if req.conversation_id else Conversation(user_id=user.id, title=mysql_safe_text(message.replace("\n", " ")[:60]), model=selected_model)
        if not req.conversation_id:
            db.add(conversation)
            db.commit()
            db.refresh(conversation)
        history = db.query(Message).filter(Message.conversation_id == conversation.id).order_by(Message.id.desc()).limit(MAX_CONTEXT_MESSAGES).all()[::-1]
        print(f"[PERF] database: {(time.perf_counter() - database_started) * 1000:.1f} ms", flush=True)
        context_started = time.perf_counter()
        sources = relevant_chunks(db, user.id, message)
        calc = calculator(message)
        if calc is not None:
            sources.append((f"Calculator result: {calc}", "calculator"))
        db.add(Message(conversation_id=conversation.id, role="user", content=mysql_safe_text(message)))
        db.commit()
        print(f"[PERF] context preparation: {(time.perf_counter() - context_started) * 1000:.1f} ms", flush=True)
        search_started = time.perf_counter()
        web_sources = await search_results(message) if req.search_enabled else []
        print(f"[PERF] web search/tool: {(time.perf_counter() - search_started) * 1000:.1f} ms", flush=True)
        sources.extend((f"{item['title']}: {item['snippet']} ({item['url']})", "web search") for item in web_sources)
        ollama_messages = [{"role": "system", "content": system_prompt(req, sources)}]
        ollama_messages += [{"role": item.role, "content": item.content} for item in history]
        ollama_messages.append({"role": "user", "content": message})
        conversation_id = conversation.id
    payload = {"model": selected_model, "messages": ollama_messages, "stream": True, "keep_alive": OLLAMA_KEEP_ALIVE}

    async def event_stream():
        answer_parts = []
        ollama_started = time.perf_counter()
        first_token_logged = False
        try:
            yield f"data: {json.dumps({'type': 'meta', 'conversation_id': conversation_id, 'sources': [{'name': name} for _, name in sources]})}\n\n"
            async for item in ollama_stream(payload):
                token = item.get("message", {}).get("content", "")
                if token:
                    if not first_token_logged:
                        print(f"[PERF] Ollama first token: {(time.perf_counter() - ollama_started) * 1000:.1f} ms", flush=True)
                        first_token_logged = True
                    answer_parts.append(token)
                    yield f"data: {json.dumps({'type': 'token', 'content': token})}\n\n"
            answer = "".join(answer_parts).strip()
            if not answer:
                raise RuntimeError("Ollama returned an empty response")
            with SessionLocal() as save_db:
                save_db.add(Message(conversation_id=conversation_id, role="assistant", content=mysql_safe_text(answer)))
                save_db.commit()
            print(f"[PERF] Ollama total: {(time.perf_counter() - ollama_started) * 1000:.1f} ms", flush=True)
            print(f"[PERF] total request: {(time.perf_counter() - request_started) * 1000:.1f} ms", flush=True)
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as exc:
            yield f"data: {json.dumps({'type': 'error', 'message': str(exc)[:500]})}\n\n"
    return StreamingResponse(event_stream(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.get("/conversations")
def conversations(user: User = Depends(current_user)):
    with SessionLocal() as db:
        rows = db.query(Conversation).filter(Conversation.user_id == user.id).order_by(Conversation.id.desc()).all()
        return [{"id": x.id, "title": x.title, "model": x.model, "created_at": x.created_at.isoformat()} for x in rows]

@app.post("/conversations")
def create_conversation(user: User = Depends(current_user)):
    with SessionLocal() as db:
        row = Conversation(user_id=user.id, title="New chat", model=OLLAMA_MODEL)
        db.add(row)
        db.commit()
        db.refresh(row)
        return {"id": row.id, "title": row.title, "model": row.model}

@app.get("/conversations/{conversation_id}")
def conversation(conversation_id: int, user: User = Depends(current_user)):
    with SessionLocal() as db:
        row = owned_conversation(db, conversation_id, user.id)
        return {"id": row.id, "title": row.title, "model": row.model, "messages": [{"id": m.id, "role": m.role, "content": m.content} for m in sorted(row.messages, key=lambda item: item.id)]}

@app.patch("/conversations/{conversation_id}")
def rename_conversation(conversation_id: int, payload: dict, user: User = Depends(current_user)):
    title = str(payload.get("title", "")).strip()[:255]
    if not title:
        raise HTTPException(400, "Title cannot be empty")
    with SessionLocal() as db:
        row = owned_conversation(db, conversation_id, user.id)
        row.title = mysql_safe_text(title)
        db.commit()
        return {"id": row.id, "title": row.title}

@app.delete("/conversations/{conversation_id}")
def delete_conversation(conversation_id: int, user: User = Depends(current_user)):
    with SessionLocal() as db:
        row = owned_conversation(db, conversation_id, user.id)
        db.delete(row)
        db.commit()
        return {"message": "deleted"}

def extract_text(suffix: str, data: bytes, path: Path) -> str:
    if suffix in {".txt", ".md", ".csv", ".json"}:
        return data.decode("utf-8", errors="ignore")
    if suffix == ".pdf":
        from pypdf import PdfReader
        return "\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages)
    if suffix == ".docx":
        from docx import Document as DocxDocument
        return "\n".join(p.text for p in DocxDocument(str(path)).paragraphs)
    return ""

@app.post("/documents")
async def upload_document(file: UploadFile = File(...), conversation_id: int | None = Form(None), user: User = Depends(current_user)):
    suffix = Path(file.filename or "").suffix.lower()
    allowed = {".txt", ".md", ".pdf", ".docx", ".csv", ".json"}
    if suffix not in allowed:
        raise HTTPException(400, "Supported files: TXT, MD, PDF, DOCX, CSV, JSON")
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File exceeds the 10 MB limit")
    safe_name = f"{uuid.uuid4().hex}{suffix}"
    path = UPLOAD_DIR / safe_name
    path.write_bytes(data)
    content = mysql_safe_text(extract_text(suffix, data, path)[:500000])
    with SessionLocal() as db:
        if conversation_id is not None:
            owned_conversation(db, conversation_id, user.id)
        doc = Document(user_id=user.id, filename=Path(file.filename or safe_name).name, content=content)
        doc.conversation_id = conversation_id
        db.add(doc)
        db.flush()
        db.add_all([DocumentChunk(document_id=doc.id, user_id=user.id, content=chunk) for chunk in split_chunks(content)])
        db.commit()
        return {"document_id": doc.id, "filename": doc.filename, "characters": len(content), "chunks": len(split_chunks(content))}

@app.get("/documents")
def documents(user: User = Depends(current_user)):
    with SessionLocal() as db:
        rows = db.query(Document).filter(Document.user_id == user.id).order_by(Document.id.desc()).all()
        return [{"id": d.id, "filename": d.filename, "characters": len(d.content)} for d in rows]

@app.get("/files/{filename}")
def files(filename: str, _: User = Depends(current_user)):
    path = UPLOAD_DIR / Path(filename).name
    if not path.exists() or path.parent != UPLOAD_DIR:
        raise HTTPException(404, "File not found")
    return FileResponse(path)

@app.get("/tool/search")
async def web_search(q: str, _: User = Depends(current_user)):
    import requests
    from bs4 import BeautifulSoup
    try:
        response = requests.post("https://html.duckduckgo.com/html/", data={"q": q[:300]}, timeout=15, headers={"User-Agent": "Yours AI/2.0"})
        soup = BeautifulSoup(response.text, "html.parser")
        return {"query": q, "results": [{"title": a.get_text(" ", strip=True), "url": a.get("href"), "snippet": (item.select_one(".result__snippet").get_text(" ", strip=True) if item.select_one(".result__snippet") else "")} for item in soup.select(".result")[:5] if (a := item.select_one(".result__a"))]}
    except requests.RequestException as exc:
        raise HTTPException(502, f"Web search unavailable: {exc}")
